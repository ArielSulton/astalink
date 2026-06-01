#!/usr/bin/env python3
"""Render proposal_2_answers.md into a styled DOCX (no pandoc needed)."""
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = Path(__file__).resolve().parent
SRC = HERE / "proposal_2_answers.md"
OUT = HERE / "AstaLink_Proposal_2.docx"

ACCENT = RGBColor(0x1F, 0x4E, 0x79)  # dark blue, matches proposal 1


def strip_inline(text: str) -> list[tuple[str, bool]]:
    """Split text into (chunk, is_bold) runs, honoring **bold** markup."""
    runs: list[tuple[str, bool]] = []
    for i, part in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if part:
            runs.append((part, i % 2 == 1))
    return runs or [(text, False)]


def add_para(doc, text, *, size=11, bold=False, color=None, space_after=6,
             align=None, style=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    for chunk, chunk_bold in strip_inline(text):
        run = p.add_run(chunk)
        run.font.size = Pt(size)
        run.font.name = "Calibri"
        run.bold = bold or chunk_bold
        if color is not None:
            run.font.color.rgb = color
    return p


def main() -> int:
    if not SRC.exists():
        print(f"ERROR: {SRC} not found", file=sys.stderr)
        return 1

    lines = SRC.read_text(encoding="utf-8").splitlines()
    doc = Document()

    # A4 + 1-inch margins
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Inches(1))

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i].rstrip()

        if not line:
            i += 1
            continue

        # Table block (GitHub pipe table)
        if line.startswith("|"):
            block = []
            while i < n and lines[i].lstrip().startswith("|"):
                block.append(lines[i].strip())
                i += 1
            rows = [
                [c.strip() for c in row.strip("|").split("|")]
                for row in block
                if not re.match(r"^\|[\s:|-]+\|?$", row)  # drop separator row
            ]
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = "Light Grid Accent 1"
                for r, row in enumerate(rows):
                    for c, cell_text in enumerate(row):
                        cell = table.cell(r, c)
                        cell.text = ""
                        para = cell.paragraphs[0]
                        for chunk, chunk_bold in strip_inline(cell_text):
                            run = para.add_run(chunk)
                            run.font.size = Pt(9.5)
                            run.font.name = "Calibri"
                            run.bold = chunk_bold or (r == 0)
                doc.add_paragraph().paragraph_format.space_after = Pt(6)
            continue

        # Headings
        if line.startswith("# "):
            add_para(doc, line[2:], size=18, bold=True, color=ACCENT,
                     space_after=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        elif line.startswith("## "):
            add_para(doc, line[3:], size=13, bold=True, color=ACCENT, space_after=4)
        elif line.startswith("### "):
            add_para(doc, line[4:], size=12, bold=True, color=ACCENT, space_after=4)
        elif line == "---":
            pass  # skip horizontal rules
        # Bullets
        elif re.match(r"^[-*] ", line):
            add_para(doc, line[2:], size=11, space_after=3, style="List Bullet")
        elif re.match(r"^\d+\. ", line):
            add_para(doc, re.sub(r"^\d+\.\s*", "", line), size=11, space_after=3,
                     style="List Number")
        else:
            add_para(doc, line, size=11)
        i += 1

    doc.save(OUT)
    print(f"Wrote {OUT}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
